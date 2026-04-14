#!/usr/bin/env python3
"""
EXECUTE MERGE: WP Word Docs (gold figures) + GitHub markdown (updated text).

Three merge strategies:
1. HIGH_OVERLAP: Section-by-section text replacement in WP doc, preserving figures
2. RESTRUCTURED_WITH_FIGS: Use WP doc as base, insert GitHub text around figures
3. NO_FIGS: Just convert GitHub markdown to docx (no figures to preserve)

Output: sapm-{slug}/paper/{slug}-merged.docx
"""

import os
import re
import shutil
import subprocess
import tempfile
from copy import deepcopy
from pathlib import Path
from lxml import etree
from docx import Document
from docx.shared import Inches, Pt, Emu
from docx.oxml.ns import qn

BASE = "/Users/erikpostnieks/Projects"
WP_DIR = f"{BASE}/deep-research/domain_papers/WP Word Docs"

WML_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
DRAW_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'
WP_NS = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
PIC_NS = 'http://schemas.openxmlformats.org/drawingml/2006/picture'

# Papers categorized by merge strategy
HIGH_OVERLAP = [
    ("Bitcoin_SAPM_v16_with_figs.docx", "bitcoin", "paper/bitcoin-current.md"),
    ("Nuclear_SAPM_v17_final (2).docx", "nuclear", "paper/nuclear-current.md"),
    ("PoS_SAPM_v9_Final_F12_edited_zingers (1).docx", "pos", "paper/pos-current.md"),
    ("DSM_SAPM_FULL_v10.docx", "deep-sea-mining", "paper/deep_sea_mining-current.md"),
]

RESTRUCTURED_WITH_FIGS = [
    ("AMR_SAPM_v10_QC.docx", "amr", "paper/amr-current.md"),
    ("Cement_Concrete_SAPM_v7_96.5 (1).docx", "cement", "paper/cement-current.md"),
    ("PST_Tobacco_v14.docx", "tobacco", "paper/tobacco-current.md"),
    ("PST_Topsoil_Erosion_v4.docx", "topsoil", "paper/topsoil-current.md"),
    ("Water_Privatization_v6_5.docx", "water-privatization", "paper/water_privatization-current.md"),
]

NO_FIGS = [
    ("CRE_Pipeline_Upgraded.docx", "cre", "paper/cre-current.md"),
    ("Gene_Drive_Propositions_K16-K22_Maximum_Strength.docx", "gene-drives", "paper/gene_drives-current.md"),
    ("PST_Big_Tech_Platform_Monopoly_Strengthened_96_5.docx", "big-tech", "paper/big_tech-current.md"),
    ("PST_Palm_Oil_REVISED_v3_MC.docx", "palm-oil", "paper/palm_oil-current.md"),
    ("PST_UPF_Upgraded (1).docx", "upf", "paper/upf-current.md"),
    ("Monoculture 96.5.docx", "monoculture", "paper/monoculture-current.md"),
    ("PST_Orbital_Debris_v3_Post_RedTeam.docx", "orbital-debris", "paper/orbital_debris-current.md"),
    ("PST_PBM_v7.docx", "pbm", "paper/pbm-current.md"),
]


def normalize_heading(text):
    """Normalize heading for matching."""
    text = text.strip()
    text = re.sub(r'^§?\s*\d+[\.\d]*[a-z]?\s*\.?\s*', '', text)
    return text.lower().strip()


def get_heading_text(para):
    """Get heading text if paragraph is a heading."""
    if para.style is None:
        return None
    if para.style.name and para.style.name.startswith('Heading'):
        return para.text.strip()
    return None


def paragraph_has_image(para):
    """Check if paragraph contains an inline image or drawing."""
    try:
        drawings = para._element.findall(f'.//{{{WML_NS}}}drawing')
        if drawings:
            return True
        # Also check for w:pict (older image format)
        picts = para._element.findall(f'.//{{{WML_NS}}}pict')
        return len(picts) > 0
    except Exception:
        return False


def get_section_boundaries(doc):
    """
    Build list of (heading_text, heading_normalized, start_idx, end_idx, has_figures).
    Each section runs from its heading to the next heading.
    """
    sections = []
    current = None

    for i, para in enumerate(doc.paragraphs):
        h = get_heading_text(para)
        if h:
            if current:
                current['end'] = i
                sections.append(current)
            current = {
                'heading': h,
                'norm': normalize_heading(h),
                'start': i,
                'end': None,
                'has_figures': False,
                'figure_indices': [],
            }
        if current and paragraph_has_image(para):
            current['has_figures'] = True
            current['figure_indices'].append(i)

    if current:
        current['end'] = len(doc.paragraphs)
        sections.append(current)

    return sections


def convert_md_to_docx(md_path, output_path):
    """Convert markdown to docx via pandoc, handling YAML issues."""
    formats = [
        'markdown-yaml_metadata_block+footnotes+subscript+superscript',
        'markdown-yaml_metadata_block',
        'markdown+footnotes+subscript+superscript',
        'markdown',
    ]
    for fmt in formats:
        cmd = ['pandoc', md_path, '-f', fmt, '-t', 'docx', '-o', output_path, '--wrap=none']
        result = subprocess.run(cmd, capture_output=True, text=True)
        if result.returncode == 0:
            return True
    return False


def copy_image_parts(src_doc, dst_doc, element):
    """
    Copy image relationship data from src_doc to dst_doc for a given XML element.
    Returns a new element with updated relationship IDs.
    """
    new_elem = deepcopy(element)

    # Find all blip references (image data)
    a_ns = 'http://schemas.openxmlformats.org/drawingml/2006/main'
    blips = new_elem.findall(f'.//{{{a_ns}}}blip')

    for blip in blips:
        embed_attr = f'{{{R_NS}}}embed'
        old_rid = blip.get(embed_attr)
        if old_rid and old_rid in src_doc.part.rels:
            rel = src_doc.part.rels[old_rid]
            if 'image' in rel.reltype:
                try:
                    # Add the image to the destination document
                    new_rid = dst_doc.part.relate_to(
                        rel.target_part,
                        rel.reltype,
                    )
                    blip.set(embed_attr, new_rid)
                except Exception as e:
                    print(f"    Warning: couldn't copy image rel {old_rid}: {e}")

    return new_elem


def merge_high_overlap(wp_path, slug, md_path, output_dir):
    """
    HIGH_OVERLAP merge: Replace text in shared sections of WP doc with GitHub text,
    preserving all figure paragraphs in their original positions.
    Add new GitHub-only sections at the end.

    Key insight: Store XML element references upfront BEFORE modifying the document,
    because wp_doc.paragraphs is dynamically computed and indices shift after removals.
    """
    print(f"  Strategy: HIGH_OVERLAP (section-by-section text replacement)")

    # Load WP doc (the base — has figures)
    wp_doc = Document(wp_path)

    # Convert GitHub md to docx
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        gh_docx_path = tmp.name
    if not convert_md_to_docx(md_path, gh_docx_path):
        print(f"  FAILED: pandoc conversion")
        return False
    gh_doc = Document(gh_docx_path)

    # Pre-collect ALL element references while indices are still valid
    # Build section map using elements, not indices
    wp_section_elems = []  # list of {heading_elem, norm, content_elems, figure_elems}
    current = None
    for para in wp_doc.paragraphs:
        h = get_heading_text(para)
        if h:
            if current:
                wp_section_elems.append(current)
            current = {
                'heading': h,
                'norm': normalize_heading(h),
                'heading_elem': para._element,
                'content_elems': [],   # non-figure paragraph elements
                'figure_elems': [],    # figure paragraph elements (PRESERVE these)
            }
        elif current:
            if paragraph_has_image(para):
                current['figure_elems'].append(para._element)
            else:
                current['content_elems'].append(para._element)
    if current:
        wp_section_elems.append(current)

    # Same for GitHub doc
    gh_section_elems = []
    current = None
    for para in gh_doc.paragraphs:
        h = get_heading_text(para)
        if h:
            if current:
                gh_section_elems.append(current)
            current = {
                'heading': h,
                'norm': normalize_heading(h),
                'heading_elem': para._element,
                'content_elems': [],
            }
        elif current:
            current['content_elems'].append(para._element)
    if current:
        gh_section_elems.append(current)

    # Build lookups
    gh_lookup = {}
    for sec in gh_section_elems:
        gh_lookup.setdefault(sec['norm'], sec)

    wp_norms = set(s['norm'] for s in wp_section_elems)

    body = wp_doc.element.body
    merged_count = 0
    figure_preserved = 0
    new_sections_added = 0

    # Phase 1: For each shared section, replace text while preserving figures
    for wp_sec in wp_section_elems:
        gh_sec = gh_lookup.get(wp_sec['norm'])
        if not gh_sec:
            continue  # WP-only section, keep as-is

        # Remove old text paragraphs (NOT figure paragraphs)
        for elem in wp_sec['content_elems']:
            try:
                body.remove(elem)
            except ValueError:
                pass

        figure_preserved += len(wp_sec['figure_elems'])

        # Insert GitHub text paragraphs after the heading element
        # If there are figures, insert text BEFORE the first figure
        if wp_sec['figure_elems']:
            # Insert new text between heading and first figure
            insert_before = wp_sec['figure_elems'][0]
            for gh_elem in reversed(gh_sec['content_elems']):
                new_para = deepcopy(gh_elem)
                insert_before.addprevious(new_para)
        else:
            # No figures — insert after heading
            insert_after = wp_sec['heading_elem']
            for gh_elem in gh_sec['content_elems']:
                new_para = deepcopy(gh_elem)
                insert_after.addnext(new_para)
                insert_after = new_para

        merged_count += 1

    # Phase 2: Add GitHub-only sections at the end
    for gh_sec in gh_section_elems:
        if gh_sec['norm'] not in wp_norms:
            # Add heading
            body.append(deepcopy(gh_sec['heading_elem']))
            # Add content
            for elem in gh_sec['content_elems']:
                body.append(deepcopy(elem))
            new_sections_added += 1

    # Save merged doc
    output_path = os.path.join(output_dir, f"{slug}-merged.docx")
    wp_doc.save(output_path)

    # Verify figure count
    merged_doc = Document(output_path)
    final_fig_count = sum(1 for p in merged_doc.paragraphs if paragraph_has_image(p))
    orig_fig_count = sum(1 for p in Document(wp_path).paragraphs if paragraph_has_image(p))

    # Word count
    merged_words = len(' '.join(p.text for p in merged_doc.paragraphs).split())

    print(f"  Merged {merged_count} sections, preserved {figure_preserved} figures, added {new_sections_added} new sections")
    print(f"  Figures: {orig_fig_count} original → {final_fig_count} in merged")
    if final_fig_count < orig_fig_count:
        print(f"  ⚠ WARNING: {orig_fig_count - final_fig_count} figures lost!")
    print(f"  Words: {merged_words:,}")
    print(f"  Output: {output_path}")

    os.unlink(gh_docx_path)
    return True


def extract_all_figures(wp_doc):
    """
    Extract ALL figures from a WP doc regardless of heading structure.
    Uses surrounding text context for placement matching.
    Returns list of figure dicts with element refs and context.
    """
    WML = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    figures = []
    paras = list(wp_doc.paragraphs)
    current_heading = None

    for i, para in enumerate(paras):
        # Track headings if they exist
        if para.style and para.style.name and para.style.name.startswith('Heading'):
            current_heading = para.text.strip()

        if not paragraph_has_image(para):
            continue

        # Gather context: preceding text (up to 5 paras back)
        context_before = []
        for j in range(max(0, i - 5), i):
            t = paras[j].text.strip()
            if t and len(t) > 10:
                context_before.append(t)

        # Following text
        context_after = []
        for j in range(i + 1, min(len(paras), i + 5)):
            t = paras[j].text.strip()
            if t and len(t) > 10:
                context_after.append(t)
                break

        # Extract section reference from surrounding text (§2.4, §5, etc.)
        section_ref = None
        for txt in context_before[-2:]:
            m = re.search(r'§(\d+[\.\d]*)', txt)
            if m:
                section_ref = m.group(0)

        figures.append({
            'heading': current_heading,
            'heading_norm': normalize_heading(current_heading) if current_heading else None,
            'para_index': i,
            'element': para._element,
            'caption': para.text.strip() if para.text.strip() else None,
            'context_before': context_before,
            'context_after': context_after,
            'section_ref': section_ref,
            'position_pct': i / max(len(paras), 1),  # fractional position in doc
        })

    return figures


def find_best_insertion_point(fig, gh_doc, gh_sections, gh_heading_lookup):
    """
    Find the best place to insert a figure in the GitHub doc.
    Tries: exact heading match → section ref match → fuzzy heading → text match → proportional position.
    """
    gh_paras = list(gh_doc.paragraphs)

    # Method 1: Exact heading match
    if fig['heading_norm'] and fig['heading_norm'] in gh_heading_lookup:
        sec = gh_heading_lookup[fig['heading_norm']]
        # Insert after first content paragraph in section
        target = min(sec['start'] + 2, sec['end'] - 1)
        if target < len(gh_paras):
            return gh_paras[target]._element, "heading"

    # Method 2: Section reference match (find §X.Y in GitHub text)
    if fig['section_ref']:
        for i, p in enumerate(gh_paras):
            if fig['section_ref'] in p.text:
                # Insert after this paragraph
                target = min(i + 1, len(gh_paras) - 1)
                return gh_paras[target]._element, "section_ref"

    # Method 3: Fuzzy heading match
    if fig['heading_norm']:
        fig_words = set(fig['heading_norm'].split()) - {'the', 'a', 'an', 'of', 'in', 'and', 'or', 'for'}
        best_match = None
        best_score = 0
        for sec in gh_sections:
            sec_words = set(sec['norm'].split()) - {'the', 'a', 'an', 'of', 'in', 'and', 'or', 'for'}
            overlap = len(fig_words & sec_words)
            if overlap > best_score and overlap >= 2:
                best_score = overlap
                best_match = sec
        if best_match:
            target = min(best_match['start'] + 2, best_match['end'] - 1)
            if target < len(gh_paras):
                return gh_paras[target]._element, "fuzzy_heading"

    # Method 4: Context text matching — find paragraphs with similar text
    for ctx_text in fig['context_before'][-2:]:
        # Use first 50 chars as search key
        key = ctx_text[:50].lower()
        for i, p in enumerate(gh_paras):
            if key in p.text.lower():
                target = min(i + 1, len(gh_paras) - 1)
                return gh_paras[target]._element, "text_match"

    # Method 5: Proportional position (last resort)
    target_idx = int(fig['position_pct'] * len(gh_paras))
    target_idx = max(0, min(target_idx, len(gh_paras) - 1))
    return gh_paras[target_idx]._element, "proportional"


def merge_restructured_with_figs(wp_path, slug, md_path, output_dir):
    """
    RESTRUCTURED merge: GitHub text is completely different structure.
    Strategy: Use GitHub text as the base, then inject WP figures at best-match positions.
    Works even for docs with no heading styles (uses text context matching).
    """
    print(f"  Strategy: RESTRUCTURED (GitHub base + inject WP figures)")

    # Extract ALL figures from WP doc (heading-independent)
    wp_doc = Document(wp_path)
    figure_info = extract_all_figures(wp_doc)
    print(f"  Extracted {len(figure_info)} figures from WP doc")

    if not figure_info:
        print(f"  No figures to inject — using GitHub text only")
        output_path = os.path.join(output_dir, f"{slug}-merged.docx")
        if not convert_md_to_docx(md_path, output_path):
            return False
        doc = Document(output_path)
        print(f"  Words: {len(' '.join(p.text for p in doc.paragraphs).split()):,}")
        print(f"  Output: {output_path}")
        return True

    # Convert GitHub md → docx as the base
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        gh_docx_path = tmp.name
    if not convert_md_to_docx(md_path, gh_docx_path):
        print(f"  FAILED: pandoc conversion")
        return False

    gh_doc = Document(gh_docx_path)
    gh_sections = get_section_boundaries(gh_doc)

    # Build heading lookup
    gh_heading_lookup = {}
    for sec in gh_sections:
        gh_heading_lookup.setdefault(sec['norm'], sec)

    body = gh_doc.element.body
    placement_methods = {}

    for fig in figure_info:
        target_elem, method = find_best_insertion_point(fig, gh_doc, gh_sections, gh_heading_lookup)
        new_fig = copy_image_parts(wp_doc, gh_doc, fig['element'])
        target_elem.addnext(new_fig)
        placement_methods[method] = placement_methods.get(method, 0) + 1

    # Save
    output_path = os.path.join(output_dir, f"{slug}-merged.docx")
    gh_doc.save(output_path)

    # Verify
    merged_doc = Document(output_path)
    final_fig_count = sum(1 for p in merged_doc.paragraphs if paragraph_has_image(p))
    merged_words = len(' '.join(p.text for p in merged_doc.paragraphs).split())

    print(f"  Placed {len(figure_info)} figures: {dict(placement_methods)}")
    print(f"  Final figure count: {final_fig_count}")
    print(f"  Words: {merged_words:,}")
    print(f"  Output: {output_path}")

    os.unlink(gh_docx_path)
    return True


def merge_no_figs(wp_path, slug, md_path, output_dir):
    """
    NO_FIGS: No figures to preserve. Just convert GitHub markdown to docx.
    WP doc is kept as reference only.
    """
    print(f"  Strategy: NO_FIGS (GitHub text only, WP has no figures)")

    output_path = os.path.join(output_dir, f"{slug}-merged.docx")
    if not convert_md_to_docx(md_path, output_path):
        print(f"  FAILED: pandoc conversion")
        return False

    doc = Document(output_path)
    words = len(' '.join(p.text for p in doc.paragraphs).split())
    print(f"  Words: {words:,}")
    print(f"  Output: {output_path}")
    return True


def main():
    print("=" * 70)
    print("EXECUTING MERGE: WP figures + GitHub text")
    print("=" * 70)

    results = {'ok': [], 'fail': []}

    # Phase 1: HIGH_OVERLAP papers
    print(f"\n{'='*70}")
    print("PHASE 1: HIGH_OVERLAP (section-by-section merge)")
    print(f"{'='*70}")
    for wp_file, slug, md_rel in HIGH_OVERLAP:
        wp_path = os.path.join(WP_DIR, wp_file)
        repo_dir = f"{BASE}/sapm-{slug}"
        md_path = os.path.join(repo_dir, md_rel)
        output_dir = os.path.join(repo_dir, "paper")
        os.makedirs(output_dir, exist_ok=True)

        print(f"\n[{slug}]")
        if not os.path.exists(wp_path):
            print(f"  SKIP: WP doc not found")
            results['fail'].append(slug)
            continue
        if not os.path.exists(md_path):
            print(f"  SKIP: GitHub md not found")
            results['fail'].append(slug)
            continue

        ok = merge_high_overlap(wp_path, slug, md_path, output_dir)
        (results['ok'] if ok else results['fail']).append(slug)

    # Phase 2: RESTRUCTURED with figures
    print(f"\n{'='*70}")
    print("PHASE 2: RESTRUCTURED WITH FIGURES (GitHub base + inject figures)")
    print(f"{'='*70}")
    for wp_file, slug, md_rel in RESTRUCTURED_WITH_FIGS:
        wp_path = os.path.join(WP_DIR, wp_file)
        repo_dir = f"{BASE}/sapm-{slug}"
        md_path = os.path.join(repo_dir, md_rel)
        output_dir = os.path.join(repo_dir, "paper")
        os.makedirs(output_dir, exist_ok=True)

        print(f"\n[{slug}]")
        if not os.path.exists(wp_path):
            print(f"  SKIP: WP doc not found")
            results['fail'].append(slug)
            continue
        if not os.path.exists(md_path):
            print(f"  SKIP: GitHub md not found")
            results['fail'].append(slug)
            continue

        ok = merge_restructured_with_figs(wp_path, slug, md_path, output_dir)
        (results['ok'] if ok else results['fail']).append(slug)

    # Phase 3: NO_FIGS papers
    print(f"\n{'='*70}")
    print("PHASE 3: NO FIGURES (GitHub text as docx)")
    print(f"{'='*70}")
    for wp_file, slug, md_rel in NO_FIGS:
        wp_path = os.path.join(WP_DIR, wp_file)
        repo_dir = f"{BASE}/sapm-{slug}"
        md_path = os.path.join(repo_dir, md_rel)
        output_dir = os.path.join(repo_dir, "paper")
        os.makedirs(output_dir, exist_ok=True)

        print(f"\n[{slug}]")
        if not os.path.exists(md_path):
            print(f"  SKIP: GitHub md not found")
            results['fail'].append(slug)
            continue

        ok = merge_no_figs(wp_path, slug, md_path, output_dir)
        (results['ok'] if ok else results['fail']).append(slug)

    # Summary
    print(f"\n{'='*70}")
    print(f"DONE: {len(results['ok'])} merged, {len(results['fail'])} failed")
    print(f"{'='*70}")
    if results['ok']:
        print(f"  OK: {', '.join(results['ok'])}")
    if results['fail']:
        print(f"  FAIL: {', '.join(results['fail'])}")

    print(f"\nOutput files: sapm-{{slug}}/paper/{{slug}}-merged.docx")
    print(f"WP originals preserved: sapm-{{slug}}/paper/{{slug}}-wp-original.docx")


if __name__ == "__main__":
    main()
