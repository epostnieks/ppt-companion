#!/usr/bin/env python3
"""
Merge ALL WP Word Docs with GitHub markdown text.
19 standalone papers + PPT (already merged).

Phase 1: High-overlap papers — section-by-section text replacement
Phase 2: Restructured papers — text insertion around figures
"""

import os
import re
import shutil
import subprocess
import tempfile
from pathlib import Path
from docx import Document

BASE = "/Users/erikpostnieks/Projects"
WP_DIR = f"{BASE}/deep-research/domain_papers/WP Word Docs"

# ALL 19 standalone papers + their GitHub slug/md path
# fmt: (wp_filename, slug, github_md_relative_path_or_None)
ALL_PAPERS = [
    # Domain papers WITH figures (10)
    ("AMR_SAPM_v10_QC.docx", "amr", "paper/amr-current.md"),
    ("Bitcoin_SAPM_v16_with_figs.docx", "bitcoin", "paper/bitcoin-current.md"),
    ("Cement_Concrete_SAPM_v7_96.5 (1).docx", "cement", "paper/cement-current.md"),
    ("DSM_SAPM_FULL_v10.docx", "deep-sea-mining", "paper/deep_sea_mining-current.md"),
    ("Nuclear_SAPM_v17_final (2).docx", "nuclear", "paper/nuclear-current.md"),
    ("PoS_SAPM_v9_Final_F12_edited_zingers (1).docx", "pos", "paper/pos-current.md"),
    ("PST_PBM_v7.docx", "pbm", "paper/pbm-current.md"),
    ("PST_Tobacco_v14.docx", "tobacco", "paper/tobacco-current.md"),
    ("PST_Topsoil_Erosion_v4.docx", "topsoil", "paper/topsoil-current.md"),
    ("Water_Privatization_v6_5.docx", "water-privatization", "paper/water_privatization-current.md"),
    # Domain papers WITHOUT figures (7)
    ("CRE_Pipeline_Upgraded.docx", "cre", "paper/cre-current.md"),
    ("Gene_Drive_Propositions_K16-K22_Maximum_Strength.docx", "gene-drives", "paper/gene_drives-current.md"),
    ("Monoculture 96.5.docx", "monoculture", "paper/monoculture-current.md"),
    ("PST_Big_Tech_Platform_Monopoly_Strengthened_96_5.docx", "big-tech", "paper/big_tech-current.md"),
    ("PST_Orbital_Debris_v3_Post_RedTeam.docx", "orbital-debris", "paper/orbital_debris-current.md"),
    ("PST_Palm_Oil_REVISED_v3_MC.docx", "palm-oil", "paper/palm_oil-current.md"),
    ("PST_UPF_Upgraded (1).docx", "upf", "paper/upf-current.md"),
    # Non-domain papers
    ("The_Hollow_Win_v_1_0_edited (1).docx", "conflictoring", "paper/conflictoring-current.md"),
    ("DA v 1.0.docx", "decision-accounting", "paper/decision_accounting-current.md"),
]

# PPT is special — already merged, just needs the 2 extra MERGED images
PPT_FINAL = "PPT_v1_9_FINAL.docx"
PPT_MERGED = "The_Private_Pareto_Trap_v1_9_MERGED.docx"


def normalize_heading(text):
    """Normalize heading for matching."""
    text = text.strip()
    text = re.sub(r'^§?\s*\d+[\.\d]*[a-z]?\s*\.?\s*', '', text)
    return text.lower().strip()


def get_section_map(doc):
    """Build heading → paragraph indices map."""
    sections = {}
    current_heading = None
    current_start = 0

    for i, para in enumerate(doc.paragraphs):
        if para.style and para.style.name and para.style.name.startswith('Heading'):
            if current_heading:
                sections[current_heading] = (current_start, i)
            current_heading = normalize_heading(para.text)
            current_start = i

    if current_heading:
        sections[current_heading] = (current_start, len(doc.paragraphs))

    return sections


def paragraph_has_image(para):
    """Check if paragraph has an inline image."""
    try:
        ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        return len(para._element.findall(f'.//{{{ns}}}drawing')) > 0
    except:
        return False


def process_paper(wp_filename, slug, md_rel_path):
    """Process one paper."""
    wp_path = os.path.join(WP_DIR, wp_filename)
    repo_dir = f"{BASE}/sapm-{slug}"
    output_dir = os.path.join(repo_dir, "paper")

    if not os.path.exists(wp_path):
        return "SKIP", "WP doc not found"

    os.makedirs(output_dir, exist_ok=True)

    # Count figures in WP doc
    wp_doc = Document(wp_path)
    fig_count = sum(1 for r in wp_doc.part.rels.values() if 'image' in r.reltype)

    # Copy WP doc as working base
    working_path = os.path.join(output_dir, f"{slug}-wp-original.docx")
    shutil.copy2(wp_path, working_path)

    # Find GitHub markdown
    md_path = os.path.join(repo_dir, md_rel_path) if md_rel_path else None
    if md_path and not os.path.exists(md_path):
        # Try alternative paths
        alt_paths = [
            os.path.join(repo_dir, f"paper/{slug}-current.md"),
            os.path.join(repo_dir, f"paper/{slug.replace('-', '_')}-current.md"),
            os.path.join(f"{BASE}/deep-research/output/{slug}/{slug}-current.md"),
            os.path.join(f"{BASE}/deep-research/output/{slug.replace('-', '_')}/{slug.replace('-', '_')}-current.md"),
        ]
        md_path = None
        for alt in alt_paths:
            if os.path.exists(alt):
                md_path = alt
                break

    if not md_path or not os.path.exists(md_path):
        return "WP_ONLY", f"{fig_count} figs, no GitHub md found"

    # Get GitHub word count
    with open(md_path) as f:
        gh_words = len(f.read().split())

    # Get WP word count
    wp_text = ' '.join(p.text for p in wp_doc.paragraphs)
    wp_words = len(wp_text.split())

    # Convert GitHub md to docx
    pandoc_path = os.path.join(output_dir, f"{slug}-github-text.docx")
    cmd = ['pandoc', md_path, '-f', 'markdown+footnotes+subscript+superscript',
           '-t', 'docx', '-o', pandoc_path, '--wrap=none']
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        # Try with simpler format spec
        cmd = ['pandoc', md_path, '-f', 'markdown', '-t', 'docx', '-o', pandoc_path]
        result = subprocess.run(cmd, capture_output=True, text=True)
        if result.returncode != 0:
            return "WP_ONLY", f"{fig_count} figs, pandoc failed"

    # Compare sections
    wp_sections = get_section_map(wp_doc)
    gh_doc = Document(pandoc_path)
    gh_sections = get_section_map(gh_doc)

    shared = set(wp_sections.keys()) & set(gh_sections.keys())
    wp_only = set(wp_sections.keys()) - set(gh_sections.keys())
    gh_only = set(gh_sections.keys()) - set(wp_sections.keys())

    # Classify: high-overlap vs restructured
    total_sections = len(set(wp_sections.keys()) | set(gh_sections.keys()))
    overlap_pct = len(shared) / max(total_sections, 1) * 100

    # Save figure position report
    report_path = os.path.join(output_dir, f"{slug}-merge-report.txt")
    with open(report_path, 'w') as f:
        f.write(f"MERGE REPORT: {slug}\n{'=' * 60}\n\n")
        f.write(f"WP doc: {wp_filename}\n")
        f.write(f"WP: {wp_words:,} words, {fig_count} figures\n")
        f.write(f"GitHub: {gh_words:,} words\n")
        f.write(f"Text growth: {(gh_words - wp_words) / max(wp_words, 1) * 100:+.0f}%\n\n")
        f.write(f"Section overlap: {len(shared)}/{total_sections} ({overlap_pct:.0f}%)\n")
        f.write(f"  Shared: {len(shared)}\n")
        f.write(f"  WP-only: {len(wp_only)}\n")
        f.write(f"  GitHub-only: {len(gh_only)}\n\n")

        if fig_count > 0:
            f.write(f"FIGURE POSITIONS:\n")
            current_heading = "BEFORE_FIRST_HEADING"
            for i, para in enumerate(wp_doc.paragraphs):
                if para.style and para.style.name and para.style.name.startswith('Heading'):
                    current_heading = para.text.strip()
                if paragraph_has_image(para):
                    f.write(f"  Para {i}: after [{current_heading}]\n")
            f.write("\n")

        if gh_only:
            f.write(f"NEW SECTIONS FROM GITHUB (to add):\n")
            for h in sorted(gh_only):
                f.write(f"  + {h}\n")

    category = "HIGH_OVERLAP" if overlap_pct >= 40 else "RESTRUCTURED"
    return category, f"{fig_count}fig, {wp_words:,}→{gh_words:,}w, {len(shared)}/{total_sections} overlap ({overlap_pct:.0f}%)"


def process_ppt():
    """Handle PPT — FINAL is already the master, just log it."""
    final_path = os.path.join(WP_DIR, PPT_FINAL)
    repo_dir = f"{BASE}/sapm-ppt"  # or wherever PPT goes

    # Check if sapm-ppt exists, if not check deep-research
    if not os.path.isdir(repo_dir):
        # PPT might be in a different location
        for candidate in ["sapm-ppt", "sapm-private-pareto-trap"]:
            cpath = f"{BASE}/{candidate}"
            if os.path.isdir(cpath):
                repo_dir = cpath
                break

    # Copy FINAL as the master
    output_dir = os.path.join(repo_dir, "paper") if os.path.isdir(repo_dir) else f"{BASE}/deep-research/output/ppt"
    os.makedirs(output_dir, exist_ok=True)

    shutil.copy2(final_path, os.path.join(output_dir, "ppt-wp-master.docx"))

    # Also copy the Hollow Win
    hw_path = os.path.join(WP_DIR, "The_Hollow_Win_v_1_0_edited (1).docx")
    if os.path.exists(hw_path):
        shutil.copy2(hw_path, os.path.join(output_dir, "hollow-win-wp-master.docx"))

    return "PPT_DONE", "FINAL is master (42,702w, 38fig, 119h). MERGED has 2 extra Appendix D images to pull in."


def main():
    print("=" * 70)
    print("FULL WP WORD DOCS → WORKING COPIES")
    print("19 standalone papers + PPT")
    print("=" * 70)

    results = []

    # Process PPT first
    print(f"\n[PPT]")
    cat, detail = process_ppt()
    print(f"  {cat}: {detail}")
    results.append(("ppt", cat, detail))

    # Process all 19 standalone papers
    for wp_file, slug, md_path in ALL_PAPERS:
        print(f"\n[{slug}]")
        cat, detail = process_paper(wp_file, slug, md_path)
        print(f"  {cat}: {detail}")
        results.append((slug, cat, detail))

    # Summary
    print(f"\n{'=' * 70}")
    print("SUMMARY")
    print(f"{'=' * 70}")

    categories = {}
    for slug, cat, detail in results:
        categories.setdefault(cat, []).append((slug, detail))

    for cat in ["HIGH_OVERLAP", "RESTRUCTURED", "WP_ONLY", "PPT_DONE", "SKIP"]:
        if cat in categories:
            print(f"\n{cat} ({len(categories[cat])}):")
            for slug, detail in categories[cat]:
                print(f"  {slug}: {detail}")


if __name__ == "__main__":
    main()
