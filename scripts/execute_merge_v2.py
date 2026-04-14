#!/usr/bin/env python3
"""
MERGE V2: Pick the best base document (most figures), update text from GitHub markdown.

For each paper:
1. Compare WP Word Doc figures vs -final.docx figures
2. Use whichever has MORE figures as the base
3. Update text from GitHub markdown (pandoc → docx → inject sections)
4. Save as {slug}-merged.docx

This preserves ALL existing figures regardless of source.
"""

import os
import re
import shutil
import subprocess
import tempfile
from copy import deepcopy
from pathlib import Path
from lxml import etree
from docx import Document

BASE = "/Users/erikpostnieks/Projects"
WP_DIR = f"{BASE}/deep-research/domain_papers/WP Word Docs"

WML_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

# All 19 standalone papers
ALL_PAPERS = [
    ("AMR_SAPM_v10_QC.docx", "amr", "paper/amr-current.md"),
    ("Bitcoin_SAPM_v16_with_figs.docx", "bitcoin", "paper/bitcoin-current.md"),
    ("Cement_Concrete_SAPM_v7_96.5 (1).docx", "cement", "paper/cement-current.md"),
    ("DSM_SAPM_FULL_v10.docx", "deep-sea-mining", "paper/deep_sea_mining-current.md"),
    ("Nuclear_SAPM_v17_final (2).docx", "nuclear", "paper/nuclear-current.md"),
    ("PoS_SAPM_v9_Final_F12_edited_zingers (1).docx", "pos", "paper/pos-current.md"),
    ("PST_PBM_v7.docx", "pbm", "paper/pbm-current.md"),
    ("PST_Tobacco_v14.docx", "tobacco", "paper/tobacco-current.md"),
    ("PST_Topsoil_Erosion_v4.docx", "topsoil", "paper/topsoil-current.md"),
    ("Water_Privatization_v6_5.docx", "water-privatization", "paper/water_privatization-current.md"),
    ("CRE_Pipeline_Upgraded.docx", "cre", "paper/cre-current.md"),
    ("Gene_Drive_Propositions_K16-K22_Maximum_Strength.docx", "gene-drives", "paper/gene_drives-current.md"),
    ("Monoculture 96.5.docx", "monoculture", "paper/monoculture-current.md"),
    ("PST_Big_Tech_Platform_Monopoly_Strengthened_96_5.docx", "big-tech", "paper/big_tech-current.md"),
    ("PST_Orbital_Debris_v3_Post_RedTeam.docx", "orbital-debris", "paper/orbital_debris-current.md"),
    ("PST_Palm_Oil_REVISED_v3_MC.docx", "palm-oil", "paper/palm_oil-current.md"),
    ("PST_UPF_Upgraded (1).docx", "upf", "paper/upf-current.md"),
]


def count_image_rels(doc):
    """Count image relationships in a document."""
    return sum(1 for r in doc.part.rels.values() if 'image' in r.reltype)


def count_words(doc):
    """Count words in a document."""
    return len(' '.join(p.text for p in doc.paragraphs).split())


def normalize_heading(text):
    text = text.strip()
    text = re.sub(r'^§?\s*\d+[\.\d]*[a-z]?\s*\.?\s*', '', text)
    return text.lower().strip()


def get_heading_text(para):
    if para.style is None:
        return None
    if para.style.name and para.style.name.startswith('Heading'):
        return para.text.strip()
    return None


def paragraph_has_image(para):
    try:
        drawings = para._element.findall(f'.//{{{WML_NS}}}drawing')
        if drawings:
            return True
        picts = para._element.findall(f'.//{{{WML_NS}}}pict')
        return len(picts) > 0
    except Exception:
        return False


def convert_md_to_docx(md_path, output_path):
    """Convert markdown to docx via pandoc."""
    formats = [
        'markdown-yaml_metadata_block+footnotes+subscript+superscript',
        'markdown-yaml_metadata_block',
        'markdown',
    ]
    for fmt in formats:
        cmd = ['pandoc', md_path, '-f', fmt, '-t', 'docx', '-o', output_path, '--wrap=none']
        result = subprocess.run(cmd, capture_output=True, text=True)
        if result.returncode == 0:
            return True
    return False


def find_best_base(wp_path, slug):
    """
    Find the document with the most figures to use as merge base.
    Returns (path, source_label, fig_count, word_count).
    """
    repo_dir = f"{BASE}/sapm-{slug}"
    paper_dir = os.path.join(repo_dir, "paper")

    candidates = []

    # WP Word Doc
    if os.path.exists(wp_path):
        try:
            doc = Document(wp_path)
            figs = count_image_rels(doc)
            words = count_words(doc)
            candidates.append((wp_path, "WP", figs, words))
        except Exception:
            pass

    # -final.docx
    for variant in [f'{slug}-final.docx', f'{slug.replace("-","_")}-final.docx']:
        p = os.path.join(paper_dir, variant)
        if os.path.exists(p):
            try:
                doc = Document(p)
                figs = count_image_rels(doc)
                words = count_words(doc)
                if words > 100:  # skip near-empty files
                    candidates.append((p, "final", figs, words))
            except Exception:
                pass
            break

    # -current.docx
    for variant in [f'{slug}-current.docx', f'{slug.replace("-","_")}-current.docx']:
        p = os.path.join(paper_dir, variant)
        if os.path.exists(p):
            try:
                doc = Document(p)
                figs = count_image_rels(doc)
                words = count_words(doc)
                if words > 100:
                    candidates.append((p, "current", figs, words))
            except Exception:
                pass
            break

    if not candidates:
        return None, "none", 0, 0

    # Pick the one with most figures; tiebreak by word count
    candidates.sort(key=lambda x: (x[2], x[3]), reverse=True)
    return candidates[0]


def merge_text_into_base(base_path, base_source, slug, md_path, output_dir):
    """
    Take the base document (with figures) and update its text from GitHub markdown.

    Strategy: Build section maps for both docs. For shared sections, replace text
    while preserving figures. For GitHub-only sections, append at end.
    """
    base_doc = Document(base_path)
    base_figs = count_image_rels(base_doc)

    # Convert GitHub md to docx for text
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        gh_docx_path = tmp.name
    if not convert_md_to_docx(md_path, gh_docx_path):
        print(f"  pandoc failed — keeping base as-is")
        shutil.copy2(base_path, os.path.join(output_dir, f"{slug}-merged.docx"))
        return True
    gh_doc = Document(gh_docx_path)

    # Build section maps using elements (not indices)
    base_sections = build_section_map(base_doc)
    gh_sections = build_section_map(gh_doc)

    # Build lookups
    gh_lookup = {s['norm']: s for s in gh_sections}
    base_norms = set(s['norm'] for s in base_sections)

    body = base_doc.element.body
    merged_count = 0
    figure_preserved = 0

    # For each base section that has a GitHub match, replace text (keep figures)
    for sec in base_sections:
        gh_sec = gh_lookup.get(sec['norm'])
        if not gh_sec:
            continue

        # Remove old text paragraphs (NOT figure paragraphs)
        for elem in sec['content_elems']:
            try:
                body.remove(elem)
            except ValueError:
                pass

        figure_preserved += len(sec['figure_elems'])

        # Insert GitHub text
        if sec['figure_elems']:
            # Insert before first figure
            insert_before = sec['figure_elems'][0]
            for gh_elem in reversed(gh_sec['content_elems']):
                insert_before.addprevious(deepcopy(gh_elem))
        else:
            # Insert after heading
            insert_after = sec['heading_elem']
            for gh_elem in gh_sec['content_elems']:
                new = deepcopy(gh_elem)
                insert_after.addnext(new)
                insert_after = new

        merged_count += 1

    # Add GitHub-only sections at end
    new_sections = 0
    for gh_sec in gh_sections:
        if gh_sec['norm'] not in base_norms:
            body.append(deepcopy(gh_sec['heading_elem']))
            for elem in gh_sec['content_elems']:
                body.append(deepcopy(elem))
            new_sections += 1

    # Save
    output_path = os.path.join(output_dir, f"{slug}-merged.docx")
    base_doc.save(output_path)

    # Verify
    merged_doc = Document(output_path)
    final_figs = count_image_rels(merged_doc)
    final_words = count_words(merged_doc)

    fig_status = ""
    if final_figs < base_figs:
        fig_status = f" ⚠ LOST {base_figs - final_figs}!"
    elif final_figs > base_figs:
        fig_status = f" (+{final_figs - base_figs} new)"

    print(f"  Merged {merged_count} sections, +{new_sections} new, {figure_preserved} figs preserved")
    print(f"  Figures: {base_figs} → {final_figs}{fig_status}")
    print(f"  Words: {final_words:,}")

    os.unlink(gh_docx_path)
    return True


def build_section_map(doc):
    """Build section map with element references (not indices)."""
    sections = []
    current = None

    for para in doc.paragraphs:
        h = get_heading_text(para)
        if h:
            if current:
                sections.append(current)
            current = {
                'heading': h,
                'norm': normalize_heading(h),
                'heading_elem': para._element,
                'content_elems': [],
                'figure_elems': [],
            }
        elif current:
            if paragraph_has_image(para):
                current['figure_elems'].append(para._element)
            else:
                current['content_elems'].append(para._element)

    if current:
        sections.append(current)
    return sections


def main():
    print("=" * 70)
    print("MERGE V2: Best base (most figures) + GitHub text")
    print("=" * 70)

    results = []

    for wp_file, slug, md_rel in ALL_PAPERS:
        wp_path = os.path.join(WP_DIR, wp_file)
        repo_dir = f"{BASE}/sapm-{slug}"
        md_path = os.path.join(repo_dir, md_rel)
        output_dir = os.path.join(repo_dir, "paper")
        os.makedirs(output_dir, exist_ok=True)

        print(f"\n[{slug}]")

        # Find best base document
        base_path, base_source, base_figs, base_words = find_best_base(wp_path, slug)
        if not base_path:
            print(f"  SKIP: no base document found")
            results.append((slug, "SKIP", 0, 0))
            continue

        print(f"  Base: {base_source} ({base_figs} figs, {base_words:,} words)")

        if not os.path.exists(md_path):
            # No GitHub markdown — just copy base as merged
            print(f"  No GitHub md — copying base as merged")
            shutil.copy2(base_path, os.path.join(output_dir, f"{slug}-merged.docx"))
            results.append((slug, base_source, base_figs, base_words))
            continue

        # Merge text into base
        ok = merge_text_into_base(base_path, base_source, slug, md_path, output_dir)
        if ok:
            merged = Document(os.path.join(output_dir, f"{slug}-merged.docx"))
            results.append((slug, base_source, count_image_rels(merged), count_words(merged)))
        else:
            results.append((slug, "FAIL", 0, 0))

    # Summary
    print(f"\n{'='*70}")
    print("SUMMARY")
    print(f"{'='*70}")
    total_figs = 0
    for slug, source, figs, words in results:
        total_figs += figs
        print(f"  {slug:25s} | base={source:7s} | {figs:3d} figs | {words:6,} words")
    print(f"\n  TOTAL: {total_figs} figures across {len(results)} papers")


if __name__ == "__main__":
    main()
